import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import os
import psutil
from matplotlib.patches import Patch
from typing import List, Dict, Tuple

# PyMCDM imports
import pymcdm
from pymcdm.methods import TOPSIS, WSM
from scipy.stats import spearmanr

# pymoo imports
from pymoo.util.nds.non_dominated_sorting import NonDominatedSorting

class DDoSDecisionAgent:
    """
    Full AI Decision Support System (MOO-MCDM-MCDA).
    Optimizes algorithm selection based on Performance, Explainability, and Resources.
    Uses precise weights defined in configuration.
    """
    
    def __init__(self, config: Dict):
        self.config = config
        self.hierarchy = config['mcdm_hierarchy']
        
        # User profiles (weights: Performance, Explainability, Resources)
        self.profiles = {
            "A": np.array([0.70, 0.15, 0.15]), # Focus Performance
            "B": np.array([0.15, 0.70, 0.15])  # Focus Explainability
        }
        # AHP profiles for sensitivity (3 dimensions)
        self.ahp_profiles = {
            "Balanced": np.array([0.33, 0.33, 0.34]),
            "Compliance-First": np.array([0.25, 0.50, 0.25]),
            "Resource-Constrained": np.array([0.30, 0.30, 0.40]),
            "Performance-First": np.array([0.50, 0.25, 0.25])
        }
        self.colors = {
            "LR": "#2E86AB",
            "DT": "#A23B72",
            "RF": "#F18F01",
            "CNN": "#C73E1D",
            "TabNet": "#3B1F2B"
        }

    def _display_name(self, name: str) -> str:
        if name.startswith("fused_"):
            return name.replace("fused_", "")
        return name

    def _display_list(self, algorithms: List[str]) -> List[str]:
        return [self._display_name(a) for a in algorithms]

    def get_system_load_penalty(self) -> float:
        """Check current local system load."""
        cpu_load = psutil.cpu_percent(interval=0.1)
        ram_load = psutil.virtual_memory().percent
        if cpu_load > 70 or ram_load > 70:
            return 1.5
        return 1.0

    def get_global_weights(self, profile_weights: np.ndarray = None) -> Dict[str, float]:
        """
        Compute precise global weights.
        W_global = W_pillar * W_local_criterion
        """
        global_weights = {}
        pillars = self.hierarchy['pillars']
        
        for i, pillar in enumerate(pillars):
            pillar_key = pillar['key']
            # Use profile weights if provided, otherwise hierarchy weights
            pillar_weight = profile_weights[i] if profile_weights is not None else pillar['weight']
            
            local_criteria = self.hierarchy['criteria'][pillar_key]
            for crit in local_criteria:
                global_weights[crit['key']] = pillar_weight * crit['weight']
                
        return global_weights

    # --- OUTILS TOPSIS + VISUALISATIONS ---

    @staticmethod
    def _build_matrix(df: pd.DataFrame) -> Tuple[np.ndarray, List[str]]:
        dims = ['dim_performance', 'dim_explainability', 'dim_resources']
        missing = [d for d in dims if d not in df.columns]
        if missing:
            raise ValueError(f"Missing dimensions in DF: {missing}")
        M = df[dims].to_numpy()
        algorithms = df['model'].tolist()
        return M, algorithms

    @staticmethod
    def _topsis_compute(M: np.ndarray, weights: np.ndarray) -> Tuple[np.ndarray, np.ndarray, np.ndarray, np.ndarray, np.ndarray, np.ndarray]:
        tau = np.array([1.0, 1.0, 1.0])
        M_norm = M / tau
        V = M_norm * weights
        # f_perf, f_expl maximize; f_res minimize
        A_plus = np.array([V[:, 0].max(), V[:, 1].max(), V[:, 2].min()])
        A_minus = np.array([V[:, 0].min(), V[:, 1].min(), V[:, 2].max()])
        D_plus = np.sqrt(((V - A_plus) ** 2).sum(axis=1))
        D_minus = np.sqrt(((V - A_minus) ** 2).sum(axis=1))
        CC = D_minus / (D_plus + D_minus + 1e-9)
        return M_norm, V, A_plus, A_minus, D_plus, D_minus, CC

    @staticmethod
    def _identify_pareto_front(M: np.ndarray) -> Tuple[List[int], List[int]]:
        n = len(M)
        is_dominated = [False] * n
        for i in range(n):
            for j in range(n):
                if i == j:
                    continue
                better_perf = M[j, 0] >= M[i, 0]
                better_expl = M[j, 1] >= M[i, 1]
                better_res = M[j, 2] <= M[i, 2]
                strictly = (M[j, 0] > M[i, 0] or M[j, 1] > M[i, 1] or M[j, 2] < M[i, 2])
                if better_perf and better_expl and better_res and strictly:
                    is_dominated[i] = True
                    break
        pareto_idx = [i for i in range(n) if not is_dominated[i]]
        dominated_idx = [i for i in range(n) if is_dominated[i]]
        return pareto_idx, dominated_idx

    # --- VISUALISATIONS ---

    def _viz_matrix_heatmap(self, M: np.ndarray, algorithms: List[str], title: str, filename: str, out_dir: str) -> str:
        fig, ax = plt.subplots(figsize=(10, 6))
        labels = self._display_list(algorithms)
        df = pd.DataFrame(M, index=labels, columns=["f_perf", "f_expl", "f_res"])
        sns.heatmap(df, annot=True, fmt=".3f", cmap="YlGnBu", linewidths=0.5, ax=ax, vmin=0, vmax=1)
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.set_xlabel("Criteria")
        ax.set_ylabel("Algorithms")
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_3d_tradeoffs(self, M: np.ndarray, algorithms: List[str], title: str, filename: str, out_dir: str) -> str:
        fig = plt.figure(figsize=(12, 9))
        ax = fig.add_subplot(111, projection='3d')
        for i, algo in enumerate(algorithms):
            algo_key = algo.replace("fused_", "")
            ax.scatter(M[i, 0], M[i, 1], M[i, 2], c=self.colors.get(algo_key, 'steelblue'), s=200, alpha=0.8, edgecolors='black')
            ax.text(M[i, 0], M[i, 1], M[i, 2], f'  {self._display_name(algo)}', fontsize=10, fontweight='bold')
        ax.set_xlabel('Performance (f_perf)')
        ax.set_ylabel('Explainability (f_expl)')
        ax.set_zlabel('Resources (f_res)')
        ax.set_title(title, fontsize=12, fontweight='bold')
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_normalization_comparison(self, M: np.ndarray, M_norm: np.ndarray, algorithms: List[str], filename: str, out_dir: str) -> str:
        fig, axes = plt.subplots(1, 2, figsize=(14, 5))
        labels = self._display_list(algorithms)
        df_M = pd.DataFrame(M, index=labels, columns=["f_perf", "f_expl", "f_res"])
        sns.heatmap(df_M, annot=True, fmt=".3f", cmap="Blues", ax=axes[0], vmin=0, vmax=1, linewidths=0.5)
        axes[0].set_title("M (raw)", fontsize=11, fontweight='bold')
        df_norm = pd.DataFrame(M_norm, index=labels, columns=["f_perf_hat", "f_expl_hat", "f_res_hat"])
        sns.heatmap(df_norm, annot=True, fmt=".3f", cmap="Greens", ax=axes[1], vmin=0, vmax=1, linewidths=0.5)
        axes[1].set_title("Normalized M", fontsize=11, fontweight='bold')
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_weighting_transformation(self, M_norm: np.ndarray, V: np.ndarray, weights: np.ndarray, profile_name: str, algorithms: List[str], filename: str, out_dir: str) -> str:
        fig, axes = plt.subplots(1, 3, figsize=(16, 5))
        labels = self._display_list(algorithms)
        df_norm = pd.DataFrame(M_norm, index=labels, columns=["f_perf_hat", "f_expl_hat", "f_res_hat"])
        sns.heatmap(df_norm, annot=True, fmt=".3f", cmap="Blues", ax=axes[0], vmin=0, vmax=1, linewidths=0.5)
        axes[0].set_title("Normalized M", fontsize=11, fontweight='bold')
        colors_w = ['steelblue', 'forestgreen', 'coral']
        bars = axes[1].bar(["w_perf", "w_expl", "w_res"], weights, color=colors_w)
        axes[1].set_ylim(0, 0.6)
        axes[1].set_title(f"Weights AHP ({profile_name})", fontsize=11, fontweight='bold')
        for bar, val in zip(bars, weights):
            axes[1].text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.02, f'{val:.2f}', ha='center', fontsize=10)
        df_V = pd.DataFrame(V, index=labels, columns=["v_perf", "v_expl", "v_res"])
        sns.heatmap(df_V, annot=True, fmt=".3f", cmap="Oranges", ax=axes[2], vmin=0, vmax=max(0.5, V.max()), linewidths=0.5)
        axes[2].set_title("V ponderee", fontsize=11, fontweight='bold')
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_profiles_comparison(self, profiles: Dict[str, np.ndarray], filename: str, out_dir: str) -> str:
        fig, ax = plt.subplots(figsize=(12, 6))
        x = np.arange(3)
        width = 0.2
        colors = ['steelblue', 'forestgreen', 'coral', 'purple']
        for i, (name, weights) in enumerate(profiles.items()):
            offset = (i - 1.5) * width
            bars = ax.bar(x + offset, weights, width, label=name, color=colors[i], alpha=0.8)
            for bar, val in zip(bars, weights):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01, f'{val:.2f}', ha='center', fontsize=8)
        ax.set_xticks(x)
        ax.set_xticklabels(["w_perf", "w_expl", "w_res"])
        ax.set_ylabel("Weights")
        ax.set_title("AHP profile comparison", fontsize=12, fontweight='bold')
        ax.legend(title="Profile")
        ax.set_ylim(0, 0.65)
        ax.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_ideal_solutions(self, V: np.ndarray, A_plus: np.ndarray, A_minus: np.ndarray, algorithms: List[str], filename: str, out_dir: str) -> str:
        fig = plt.figure(figsize=(12, 9))
        ax = fig.add_subplot(111, projection='3d')
        for i, algo in enumerate(algorithms):
            algo_key = algo.replace("fused_", "")
            ax.scatter(V[i, 0], V[i, 1], V[i, 2], c=self.colors.get(algo_key, 'steelblue'), s=200, alpha=0.7)
            ax.text(V[i, 0], V[i, 1], V[i, 2], f'  {self._display_name(algo)}', fontsize=10)
        ax.scatter(*A_plus, c='green', s=400, marker='*', edgecolors='black', linewidths=2, label='A_plus')
        ax.scatter(*A_minus, c='red', s=400, marker='X', edgecolors='black', linewidths=2, label='A_minus')
        ax.set_xlabel('v_perf')
        ax.set_ylabel('v_expl')
        ax.set_zlabel('v_res')
        ax.set_title('Ideal and anti-ideal solutions', fontsize=12, fontweight='bold')
        ax.legend(loc='upper left')
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_distances(self, D_plus: np.ndarray, D_minus: np.ndarray, algorithms: List[str], filename: str, out_dir: str) -> str:
        fig, ax = plt.subplots(figsize=(10, 6))
        x = np.arange(len(algorithms))
        width = 0.35
        bars1 = ax.bar(x - width/2, D_plus, width, label='D_plus', color='green', alpha=0.7)
        bars2 = ax.bar(x + width/2, D_minus, width, label='D_minus', color='red', alpha=0.7)
        ax.set_xticks(x)
        ax.set_xticklabels(self._display_list(algorithms))
        ax.set_ylabel("Distance")
        ax.set_title("TOPSIS distances", fontsize=12, fontweight='bold')
        ax.legend()
        for bar in list(bars1) + list(bars2):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.005, f'{bar.get_height():.3f}', ha='center', fontsize=8)
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_closeness(self, CC: np.ndarray, algorithms: List[str], profile_name: str, filename: str, out_dir: str) -> str:
        sorted_idx = np.argsort(CC)[::-1]
        sorted_cc = CC[sorted_idx]
        sorted_labels = [self._display_name(algorithms[i]) for i in sorted_idx]
        fig, ax = plt.subplots(figsize=(10, 6))
        colors = plt.cm.RdYlGn(sorted_cc)
        bars = ax.barh(sorted_labels, sorted_cc, color=colors, edgecolor='black')
        ax.set_xlim(0, 1)
        ax.set_xlabel("CC")
        ax.set_title(f"TOPSIS ranking - {profile_name}", fontsize=12, fontweight='bold')
        for i, (bar, cc) in enumerate(zip(bars, sorted_cc)):
            ax.text(cc + 0.02, bar.get_y() + bar.get_height()/2, f'#{i+1} CC={cc:.3f}', va='center', fontsize=9)
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_radar(self, M: np.ndarray, algorithms: List[str], filename: str, out_dir: str) -> str:
        categories = ['Performance', 'Explainability', 'Resources(inv)']
        N = len(categories)
        angles = [n / float(N) * 2 * np.pi for n in range(N)]
        angles += angles[:1]
        fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(polar=True))
        for i, algo in enumerate(algorithms):
            values = [M[i, 0], M[i, 1], 1 - M[i, 2]]
            values += values[:1]
            algo_key = algo.replace("fused_", "")
            color = self.colors.get(algo_key, f'C{i}')
            ax.plot(angles, values, linewidth=2, linestyle='solid', label=self._display_name(algo), color=color)
            ax.fill(angles, values, alpha=0.1, color=color)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(categories, fontsize=11)
        ax.set_ylim(0, 1)
        ax.set_title("Comparative radar", fontsize=12, fontweight='bold', pad=20)
        ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1))
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_2d_perf_expl(self, M: np.ndarray, CC: np.ndarray, algorithms: List[str], filename: str, out_dir: str) -> str:
        fig, ax = plt.subplots(figsize=(10, 8))
        scatter = ax.scatter(M[:, 0], M[:, 1], s=CC * 800, c=M[:, 2], cmap='RdYlGn_r', alpha=0.7, edgecolors='black', linewidths=2)
        for i, algo in enumerate(algorithms):
            ax.annotate(f'{self._display_name(algo)}\\nCC={CC[i]:.3f}', (M[i, 0], M[i, 1]), textcoords="offset points", xytext=(10, 10), fontsize=10)
        ax.set_xlabel('f_perf')
        ax.set_ylabel('f_expl')
        ax.set_title('Performance vs Explainability', fontsize=12, fontweight='bold')
        cbar = plt.colorbar(scatter)
        cbar.set_label('f_res', fontsize=10)
        ax.grid(True, alpha=0.3)
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_pareto_3d(self, M: np.ndarray, algorithms: List[str], pareto_idx: List[int], dominated_idx: List[int], filename: str, out_dir: str) -> str:
        fig = plt.figure(figsize=(14, 10))
        ax = fig.add_subplot(111, projection='3d')
        for i in dominated_idx:
            ax.scatter(M[i, 0], M[i, 1], M[i, 2], c='gray', s=150, alpha=0.4, marker='o')
            ax.text(M[i, 0], M[i, 1], M[i, 2], f'  {self._display_name(algorithms[i])}', fontsize=9, color='gray')
        for i in pareto_idx:
            algo_key = algorithms[i].replace("fused_", "")
            ax.scatter(M[i, 0], M[i, 1], M[i, 2], c=self.colors.get(algo_key, 'green'), s=400, alpha=0.9, edgecolors='black', linewidths=2, marker='*')
            ax.text(M[i, 0], M[i, 1], M[i, 2], f'  {self._display_name(algorithms[i])} *', fontsize=11, fontweight='bold')
        ax.set_xlabel('Performance')
        ax.set_ylabel('Explainability')
        ax.set_zlabel('Resources')
        ax.set_title('Pareto front', fontsize=12, fontweight='bold')
        legend_elements = [
            Patch(facecolor='green', label=f'Pareto-optimal ({len(pareto_idx)})'),
            Patch(facecolor='gray', label=f'Dominated ({len(dominated_idx)})')
        ]
        ax.legend(handles=legend_elements, loc='upper left')
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_pareto_2d(self, M: np.ndarray, algorithms: List[str], pareto_idx: List[int], dominated_idx: List[int], filename: str, out_dir: str) -> str:
        fig, axes = plt.subplots(1, 3, figsize=(16, 5))
        projections = [
            (0, 1, 'Performance', 'Explainability'),
            (0, 2, 'Performance', 'Resources(inv)'),
            (1, 2, 'Explainability', 'Resources(inv)')
        ]
        for ax, (xi, yi, xlabel, ylabel) in zip(axes, projections):
            for i in dominated_idx:
                y_val = 1 - M[i, yi] if yi == 2 else M[i, yi]
                ax.scatter(M[i, xi], y_val, c='gray', s=100, alpha=0.4)
                ax.annotate(self._display_name(algorithms[i]), (M[i, xi], y_val), fontsize=8, color='gray')
            for i in pareto_idx:
                y_val = 1 - M[i, yi] if yi == 2 else M[i, yi]
                algo_key = algorithms[i].replace("fused_", "")
                ax.scatter(M[i, xi], y_val, c=self.colors.get(algo_key, 'green'), s=200, marker='*', edgecolors='black', linewidths=1)
                ax.annotate(f'{self._display_name(algorithms[i])} *', (M[i, xi], y_val), fontsize=9, fontweight='bold')
            ax.set_xlabel(xlabel, fontsize=10)
            ax.set_ylabel(ylabel, fontsize=10)
            ax.set_title(f'{xlabel} vs {ylabel}', fontsize=11, fontweight='bold')
            ax.grid(True, alpha=0.3)
        plt.suptitle('2D Pareto projections', fontsize=14, fontweight='bold')
        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_pareto_list(self, algorithms: List[str], pareto_idx: List[int], dominated_idx: List[int], filename: str, out_dir: str) -> str:
        pareto_list = [self._display_name(algorithms[i]) for i in pareto_idx]
        dominated_list = [self._display_name(algorithms[i]) for i in dominated_idx]
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.axis('off')
        ax.set_title('Pareto solutions list', fontsize=12, fontweight='bold')

        left_text = "Non-dominated (Pareto):\\n" + "\\n".join(pareto_list) if pareto_list else "Non-dominated (Pareto):\\nNone"
        right_text = "Dominated:\\n" + "\\n".join(dominated_list) if dominated_list else "Dominated:\\nNone"

        ax.text(0.02, 0.95, left_text, va='top', ha='left', fontsize=10, family='monospace')
        ax.text(0.52, 0.95, right_text, va='top', ha='left', fontsize=10, family='monospace')

        plt.tight_layout()
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_threshold_variation(
        self,
        scores: Dict[str, float],
        threshold: float,
        direction: str,
        title: str,
        formula: str,
        filename: str,
        out_dir: str,
    ) -> Tuple[str, List[str], List[str]]:
        items = [(self._display_name(k), v) for k, v in scores.items()]
        reverse = True if direction == "gte" else False
        items_sorted = sorted(items, key=lambda x: x[1], reverse=reverse)

        admissible = []
        rejected = []
        for algo, score in items_sorted:
            if direction == "gte":
                (admissible if score >= threshold else rejected).append(algo)
            else:
                (admissible if score <= threshold else rejected).append(algo)

        fig, axes = plt.subplots(1, 2, figsize=(13, 6))
        axes[0].axis('off')
        axes[0].set_title("Solutions", fontsize=11, fontweight='bold')
        left_text = "Admissible:\\n" + "\\n".join(admissible) if admissible else "Admissible:\\nNone"
        right_text = "Rejected:\\n" + "\\n".join(rejected) if rejected else "Rejected:\\nNone"
        axes[0].text(0.02, 0.95, left_text, va='top', ha='left', fontsize=9, family='monospace')
        axes[0].text(0.52, 0.95, right_text, va='top', ha='left', fontsize=9, family='monospace')

        labels = [algo for algo, _ in items_sorted]
        values = [score for _, score in items_sorted]
        colors = ["#2ecc71" if algo in admissible else "#bdc3c7" for algo in labels]
        axes[1].barh(labels, values, color=colors, alpha=0.9)
        axes[1].axvline(threshold, color='red', linestyle='--', linewidth=1.5, label=f"Threshold {threshold:.1f}")
        axes[1].set_xlabel("Score")
        axes[1].set_title(title, fontsize=11, fontweight='bold')
        axes[1].legend(loc="lower right", fontsize=8)
        axes[1].grid(axis='x', alpha=0.3)

        fig.suptitle(f"{title} - Threshold {threshold:.1f}", fontsize=12, fontweight='bold')
        fig.text(0.5, 0.02, f"Formula: {formula}", ha='center', fontsize=9, style='italic')
        plt.tight_layout()
        os.makedirs(out_dir, exist_ok=True)
        path = os.path.join(out_dir, filename)
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path, admissible, rejected

    def _write_variation_docx(
        self,
        docx_path: str,
        title: str,
        threshold: float,
        formula: str,
        admissible: List[str],
        rejected: List[str],
        criteria_summary: str,
    ) -> str:
        total = len(admissible) + len(rejected)
        if total == 0:
            total = 1
        if len(admissible) == 0:
            interp = "Threshold tres strict: no admissible solution. Relax the constraint or adjust the weights."
        elif len(admissible) == total:
            interp = "Threshold permissif: all solutions are admissible. The constraint does not discriminate."
        else:
            interp = (
                "Threshold intermediaire: la contrainte filtre une partie des solutions. "
                "This illustrates the trade-off between performance, explainability, and resources."
            )

        try:
            from docx import Document
            doc = Document()
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"Threshold applique: {threshold:.1f}")
            doc.add_paragraph(f"Formula: {formula}")
            doc.add_paragraph(f"Admissible solutions: {len(admissible)}/{total}")
            doc.add_heading("Admissible solutions", level=2)
            for algo in admissible:
                doc.add_paragraph(algo, style='List Bullet')
            doc.add_heading("Rejected solutions", level=2)
            for algo in rejected:
                doc.add_paragraph(algo, style='List Bullet')
            doc.add_heading("Interpretation", level=2)
            doc.add_paragraph(interp)
            doc.add_heading("Link to assessment criteria", level=2)
            doc.add_paragraph(criteria_summary)
            doc.save(docx_path)
            return docx_path
        except Exception:
            md_path = docx_path.replace(".docx", ".md")
            with open(md_path, "w") as f:
                f.write(f"# {title}\n\n")
                f.write(f"- Threshold applique: {threshold:.1f}\n")
                f.write(f"- Formula: {formula}\n")
                f.write(f"- Admissible solutions: {len(admissible)}/{total}\n\n")
                f.write("## Admissible solutions\n")
                for algo in admissible:
                    f.write(f"- {algo}\n")
                f.write("\n## Rejected solutions\n")
                for algo in rejected:
                    f.write(f"- {algo}\n")
                f.write("\n## Interpretation\n")
                f.write(f"{interp}\n\n")
                f.write("## Link to assessment criteria\n")
                f.write(criteria_summary + "\n")
            return md_path

    def _generate_threshold_variations(
        self,
        df: pd.DataFrame,
        score_series: pd.Series,
        thresholds: List[float],
        direction: str,
        title_prefix: str,
        formula: str,
        out_dir: str,
        report_dir: str,
    ) -> List[str]:
        os.makedirs(out_dir, exist_ok=True)
        os.makedirs(report_dir, exist_ok=True)
        results = []
        files = []
        scores = {row["model"]: float(score_series.loc[idx]) for idx, row in df.iterrows()}
        weights_summary = []
        for pillar in self.hierarchy["pillars"]:
            weights_summary.append(f"{pillar['name']}: {pillar['weight']:.2f}")
        criteria_summary = "Weights des piliers: " + ", ".join(weights_summary)

        for threshold in thresholds:
            slug = title_prefix.lower().replace(" ", "_")
            filename = f"{slug}_threshold_{int(threshold)}.png"
            plot_path, admissible, rejected = self._viz_threshold_variation(
                scores,
                threshold,
                direction,
                title_prefix,
                formula,
                filename,
                out_dir,
            )
            files.append(plot_path)
            results.append({"threshold": threshold, "admissible": len(admissible)})
            docx_path = os.path.join(report_dir, f"{slug}_threshold_{int(threshold)}.docx")
            self._write_variation_docx(
                docx_path,
                f"{title_prefix} - Threshold {threshold:.1f}",
                threshold,
                formula,
                admissible,
                rejected,
                criteria_summary,
            )

        # Resume des admissibles par seuil
        if results:
            fig, ax = plt.subplots(figsize=(8, 4))
            xs = [r["threshold"] for r in results]
            ys = [r["admissible"] for r in results]
            ax.plot(xs, ys, marker='o')
            ax.set_title(f"Number of admissible solutions - {title_prefix}", fontsize=11, fontweight='bold')
            ax.set_xlabel("Threshold")
            ax.set_ylabel("Nombre admissibles")
            ax.grid(alpha=0.3)
            summary_path = os.path.join(out_dir, f"{title_prefix.lower().replace(' ', '_')}_resume.png")
            plt.tight_layout()
            plt.savefig(summary_path, dpi=150, bbox_inches='tight')
            plt.close()
            files.append(summary_path)
        return files

    def _viz_sensitivity_weights(self, M: np.ndarray, algorithms: List[str], out_dir: str) -> str:
        results = {}
        for profile_name, weights in self.ahp_profiles.items():
            _, _, _, _, _, _, cc = self._topsis_compute(M, weights)
            results[profile_name] = cc
        fig, ax = plt.subplots(figsize=(12, 6))
        labels = self._display_list(algorithms)
        df = pd.DataFrame(results, index=labels)
        df.plot(kind='bar', ax=ax, width=0.8, alpha=0.8)
        ax.set_ylabel("CC")
        ax.set_xlabel("Algorithm")
        ax.set_title("Sensitivity to weights", fontsize=12, fontweight='bold')
        ax.legend(title="Profile", loc='upper right')
        ax.set_xticklabels(labels, rotation=0)
        ax.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        path = os.path.join(out_dir, "viz_sensitivity_weights.png")
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_sensitivity_criteria(self, M: np.ndarray, algorithms: List[str], weights: np.ndarray, out_dir: str) -> str:
        results = {'Full Model': self._topsis_compute(M, weights)[-1]}
        criteria = ['f_perf', 'f_expl', 'f_res']
        for i, criterion in enumerate(criteria):
            mask = [j for j in range(3) if j != i]
            M_reduced = M[:, mask]
            w_reduced = weights[mask]
            w_reduced = w_reduced / w_reduced.sum()
            tau = np.array([1.0, 1.0])
            M_norm = M_reduced / tau
            V = M_norm * w_reduced
            if i == 2:
                A_plus = V.max(axis=0)
                A_minus = V.min(axis=0)
            else:
                # keep last as cost if present
                A_plus = np.array([V[:, 0].max(), V[:, 1].min() if mask[1] == 2 else V[:, 1].max()])
                A_minus = np.array([V[:, 0].min(), V[:, 1].max() if mask[1] == 2 else V[:, 1].min()])
            D_plus = np.sqrt(((V - A_plus) ** 2).sum(axis=1))
            D_minus = np.sqrt(((V - A_minus) ** 2).sum(axis=1))
            CC = D_minus / (D_plus + D_minus + 1e-9)
            results[f'Sans {criterion}'] = CC
        fig, ax = plt.subplots(figsize=(12, 6))
        labels = self._display_list(algorithms)
        df = pd.DataFrame(results, index=labels)
        x = np.arange(len(algorithms))
        width = 0.2
        colors = ['green', 'steelblue', 'orange', 'red']
        for i, (scenario, values) in enumerate(results.items()):
            offset = (i - len(results)/2) * width
            ax.bar(x + offset, values, width, label=scenario, color=colors[i % len(colors)], alpha=0.8)
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_ylabel("CC")
        ax.set_title("Sensitivity criteria removal", fontsize=12, fontweight='bold')
        ax.legend()
        ax.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        path = os.path.join(out_dir, "viz_sensitivity_criteria.png")
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_sensitivity_algorithms(self, M: np.ndarray, algorithms: List[str], weights: np.ndarray, out_dir: str) -> str:
        results = {}
        ranking_full = self._topsis_compute(M, weights)[-1]
        results['Complet'] = {algo: rank for algo, rank in zip(algorithms, np.argsort(np.argsort(-ranking_full)) + 1)}
        for i, removed_algo in enumerate(algorithms):
            mask = [j for j in range(len(algorithms)) if j != i]
            M_reduced = M[mask]
            algos_reduced = [algorithms[j] for j in mask]
            cc = self._topsis_compute(M_reduced, weights)[-1]
            ranks = np.argsort(np.argsort(-cc)) + 1
            results[f'Sans {removed_algo}'] = {algo: ranks[k] for k, algo in enumerate(algos_reduced)}
        fig, ax = plt.subplots(figsize=(12, 6))
        scenarios = list(results.keys())
        x = np.arange(len(algorithms))
        width = 0.15
        colors = plt.cm.tab10(np.linspace(0, 1, len(scenarios)))
        labels = self._display_list(algorithms)
        for i, (scenario, ranks) in enumerate(results.items()):
            offset = (i - len(scenarios)/2) * width
            values = [ranks.get(algo, 0) for algo in algorithms]
            ax.bar(x + offset, values, width, label=scenario, color=colors[i], alpha=0.8)
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_ylabel("Rang (1=meilleur)")
        ax.set_title("Sensitivity algorithm removal", fontsize=12, fontweight='bold')
        ax.legend(loc='upper right', fontsize=8)
        ax.invert_yaxis()
        ax.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        path = os.path.join(out_dir, "viz_sensitivity_algorithms.png")
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _viz_constraints_impact(self, df: pd.DataFrame, algorithms: List[str], out_dir: str) -> str:
        profiles = {
            'micro_enterprise': {'tau_memory': 256 * 1024 * 1024, 'tau_cpu': 50.0},
            'small_enterprise': {'tau_memory': 512 * 1024 * 1024, 'tau_cpu': 70.0},
            'medium_enterprise': {'tau_memory': 1024 * 1024 * 1024, 'tau_cpu': 80.0},
            'large_enterprise': {'tau_memory': 4 * 1024 * 1024 * 1024, 'tau_cpu': 95.0},
        }
        results = {}
        for profile_name, config in profiles.items():
            f_res_vals = []
            for _, row in df.iterrows():
                mem_bytes = float(row.get('memory_bytes', config['tau_memory']))
                cpu_percent = float(row.get('cpu_percent', config['tau_cpu']))
                m_norm = mem_bytes / config['tau_memory']
                cpu_norm = cpu_percent / config['tau_cpu']
                f_res = 0.5 * m_norm + 0.5 * cpu_norm
                f_res_vals.append(f_res)
            admissible = sum(1 for f in f_res_vals if f <= 1.0)
            results[profile_name] = {'f_res_values': f_res_vals, 'admissible_count': admissible, 'total': len(algorithms)}

        fig, axes = plt.subplots(1, 2, figsize=(14, 5))
        profiles_list = list(results.keys())
        x = np.arange(len(algorithms))
        width = 0.2
        colors = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D']
        labels = self._display_list(algorithms)
        for i, profile in enumerate(profiles_list):
            offset = (i - 1.5) * width
            axes[0].bar(x + offset, results[profile]['f_res_values'], width, label=profile.replace('_', ' ').title(), color=colors[i], alpha=0.8)
        axes[0].axhline(y=1.0, color='red', linestyle='--', linewidth=2, label='Threshold admissibilite')
        axes[0].set_xticks(x)
        axes[0].set_xticklabels(labels, rotation=45, ha='right')
        axes[0].set_ylabel('f_res')
        axes[0].set_title('f_res par profil', fontsize=12, fontweight='bold')
        axes[0].legend(loc='upper left', fontsize=8)
        admissible_counts = [results[p]['admissible_count'] for p in profiles_list]
        totals = [results[p]['total'] for p in profiles_list]
        bars = axes[1].bar(profiles_list, admissible_counts, color=colors, alpha=0.8)
        axes[1].axhline(y=totals[0], color='gray', linestyle='--', alpha=0.5)
        for bar, count, total in zip(bars, admissible_counts, totals):
            axes[1].text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, f'{count}/{total}', ha='center', fontsize=10, fontweight='bold')
        axes[1].set_ylabel('Nombre admissibles')
        axes[1].set_title('Admissibility by SME profile', fontsize=12, fontweight='bold')
        axes[1].set_xticklabels([p.replace('_', '\\n').title() for p in profiles_list], rotation=0)
        plt.tight_layout()
        path = os.path.join(out_dir, "viz_constraints_impact.png")
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    # --- BLOCK 1: MOO (Generation and Filtering) ---
    
    def run_moo_phase(self, results_df: pd.DataFrame) -> pd.DataFrame:
        """MOO phase: aggregation and Pareto front."""
        print("\n>>> START OF MOO PHASE (Multi-Objective Optimization)")
        df = results_df.copy()
        penalty = self.get_system_load_penalty()
        
        for pillar in self.hierarchy['pillars']:
            key = pillar['key']
            criteria_keys = [c['key'] for c in self.hierarchy['criteria'][key]]
            
            # Normalisation Min-Max pour la moyenne
            temp_df = df[criteria_keys].copy()
            for c_key in criteria_keys:
                c_info = next(item for item in self.hierarchy['criteria'][key] if item["key"] == c_key)
                if c_info['type'] == 'benefit':
                    temp_df[c_key] = (df[c_key] - df[c_key].min()) / (df[c_key].max() - df[c_key].min() + 1e-9)
                else:
                    temp_df[c_key] = (df[c_key].max() - df[c_key]) / (df[c_key].max() - df[c_key].min() + 1e-9)
            
            df[f'dim_{key}'] = temp_df.mean(axis=1)
            if key == 'resources' and penalty > 1.0:
                df[f'dim_{key}'] *= (1.0 / penalty)

        objectives = ['dim_performance', 'dim_explainability', 'dim_resources']
        matrix = df[objectives].to_numpy()
        nds = NonDominatedSorting().do(-matrix)
        df['is_pareto_efficient'] = False
        df.loc[df.index[nds[0]], 'is_pareto_efficient'] = True
        return df

    # --- BLOCK 2: MCDM (Decision Modeling) ---

    def run_mcdm_phase(self, df: pd.DataFrame, profile_name: str) -> pd.DataFrame:
        """MCDM phase: ranking via TOPSIS and WSM."""
        print(f"\n>>> START OF MCDM PHASE (Profile {profile_name})")
        weights = self.profiles[profile_name]
        objectives = ['dim_performance', 'dim_explainability', 'dim_resources']
        matrix = df[objectives].to_numpy()
        types = np.array([1, 1, 1])
        
        df[f'score_topsis_{profile_name}'] = TOPSIS()(matrix, weights, types)
        df[f'score_wsm_{profile_name}'] = (matrix * weights).sum(axis=1)
        return df

    # --- BLOCK 3: MCDA (Analysis and Justification) ---

    def run_mcda_analysis(self, df: pd.DataFrame, profile_name: str) -> Dict:
        """MCDA phase: sensitivity analysis and correlation."""
        print(f"\n>>> START OF MCDA PHASE (Profile {profile_name})")
        original_weights = self.profiles[profile_name].copy()
        idx_main = 0 if profile_name == "A" else 1
        original_winner = df.sort_values(by=f'score_topsis_{profile_name}', ascending=False).iloc[0]['model']
        
        test_weights = original_weights.copy()
        test_weights[idx_main] *= 1.05
        test_weights /= test_weights.sum()
        
        objectives = ['dim_performance', 'dim_explainability', 'dim_resources']
        new_scores = TOPSIS()(df[objectives].to_numpy(), test_weights, np.array([1, 1, 1]))
        new_winner = df.iloc[np.argmax(new_scores)]['model']
        
        corr, _ = spearmanr(df[f'score_topsis_{profile_name}'], df[f'score_wsm_{profile_name}'])
        return {"is_stable": original_winner == new_winner, "spearman_corr": corr, "winner": original_winner}

    def rank_models(self, results_df: pd.DataFrame) -> pd.DataFrame:
        """Wrapper to keep pipeline compatibility."""
        df = self.run_moo_phase(results_df)
        df = self.run_mcdm_phase(df, "A")
        return df

    def visualize_sad(self, results_df: pd.DataFrame, output_dir: str):
        """Generates all MCDM/MOO visualizations."""
        os.makedirs(output_dir, exist_ok=True)
        plt.style.use('ggplot')
        
        df = self.run_moo_phase(results_df)
        df = self.run_mcdm_phase(df, "A")
        df = self.run_mcdm_phase(df, "B")
        df = df.copy()
        df["model_label"] = df["model"].apply(self._display_name)
        
        all_crit_info = [c for p in self.hierarchy['criteria'].values() for c in p]
        criteria_keys = [c['key'] for c in all_crit_info]
        norm_df = self.normalize_matrix(results_df, all_crit_info)

        # 1. Heatmap
        plt.figure(figsize=(12, 8))
        norm_df = norm_df.copy()
        norm_df["model_label"] = norm_df["model"].apply(self._display_name)
        sns.heatmap(norm_df.set_index('model_label')[criteria_keys], annot=True, cmap='YlGnBu', fmt='.2f')
        plt.title('Normalized decision matrix')
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, 'decision_matrix_heatmap.png'))
        plt.close()

        # 2. Unified chart
        fig, axes = plt.subplots(1, 3, figsize=(18, 6))
        dims = ['dim_performance', 'dim_explainability', 'dim_resources']
        titles = ['Performance', 'Explainability', 'Resources']
        colors = ['#3498db', '#e67e22', '#2ecc71']
        for i, dim in enumerate(dims):
            df_sorted = df.sort_values(by=dim, ascending=True)
            axes[i].barh(df_sorted['model_label'], df_sorted[dim], color=colors[i])
            axes[i].set_title(titles[i])
        plt.suptitle('Unified dimension comparison')
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, 'unified_dimensions_comparison.png'))
        plt.close()

        # 3. Scatter 3D
        fig = plt.figure(figsize=(10, 8))
        ax = fig.add_subplot(111, projection='3d')
        for i, row in df.iterrows():
            color = 'green' if row['is_pareto_efficient'] else 'red'
            ax.scatter(row['dim_performance'], row['dim_explainability'], row['dim_resources'], c=color, s=100)
            ax.text(row['dim_performance'], row['dim_explainability'], row['dim_resources'], row['model_label'])
        ax.set_xlabel('Performance')
        ax.set_ylabel('Explainability')
        ax.set_zlabel('Resources')
        plt.savefig(os.path.join(output_dir, '3d_solution_space.png'))
        plt.close()

        # 4. Radar
        winner_a = df.sort_values(by='score_topsis_A', ascending=False).iloc[0]
        angles = np.linspace(0, 2*np.pi, len(dims), endpoint=False).tolist()
        angles += angles[:1]
        fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
        for i, row in df.iterrows():
            values = row[dims].tolist()
            values += values[:1]
            ax.plot(angles, values, label=row['model_label'], linewidth=4 if row['model'] == winner_a['model'] else 1)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(['Performance', 'Explainability', 'Resources'])
        plt.legend(loc='upper right', bbox_to_anchor=(0.1, 0.1))
        plt.savefig(os.path.join(output_dir, 'radar_comparison.png'))
        plt.close()

        # ---- Advanced TOPSIS visualizations ----
        M, algorithms = self._build_matrix(df)
        default_profile = "Balanced"
        weights = self.ahp_profiles.get(default_profile, np.array([0.33, 0.33, 0.34]))
        M_norm, V, A_plus, A_minus, D_plus, D_minus, CC = self._topsis_compute(M, weights)
        pareto_idx, dominated_idx = self._identify_pareto_front(M)

        self._viz_matrix_heatmap(M, algorithms, "Step 0: Decision matrix M", "viz_01_matrix_M.png", output_dir)
        self._viz_3d_tradeoffs(M, algorithms, "3D trade-off space", "viz_02_3d_tradeoffs.png", output_dir)
        self._viz_normalization_comparison(M, M_norm, algorithms, "viz_03_normalization.png", output_dir)
        self._viz_weighting_transformation(M_norm, V, weights, default_profile, algorithms, "viz_04_weighting.png", output_dir)
        self._viz_profiles_comparison(self.ahp_profiles, "viz_05_profiles.png", output_dir)
        self._viz_ideal_solutions(V, A_plus, A_minus, algorithms, "viz_06_ideal_solutions.png", output_dir)
        self._viz_distances(D_plus, D_minus, algorithms, "viz_07_distances.png", output_dir)
        self._viz_closeness(CC, algorithms, default_profile, "viz_08_closeness.png", output_dir)
        self._viz_radar(M, algorithms, "viz_09_radar.png", output_dir)
        self._viz_2d_perf_expl(M, CC, algorithms, "viz_10_2d_perf_expl.png", output_dir)
        self._viz_pareto_3d(M, algorithms, pareto_idx, dominated_idx, "viz_11_pareto_3d.png", output_dir)
        self._viz_pareto_2d(M, algorithms, pareto_idx, dominated_idx, "viz_12_pareto_2d.png", output_dir)
        self._viz_pareto_list(algorithms, pareto_idx, dominated_idx, "viz_13_pareto_list.png", output_dir)
        self._viz_sensitivity_weights(M, algorithms, output_dir)
        self._viz_sensitivity_criteria(M, algorithms, weights, output_dir)
        self._viz_sensitivity_algorithms(M, algorithms, weights, output_dir)
        self._viz_constraints_impact(df, algorithms, output_dir)

        # ---- Constraint variations (100 -> 30) ----
        thresholds = list(range(100, 20, -10))
        variations_root = os.path.join(output_dir, "variations")

        # Resources (CPU/RAM mean in %)
        if "cpu_percent" in df.columns and "ram_percent" in df.columns:
            resource_scores = ((df["cpu_percent"].fillna(100.0) + df["ram_percent"].fillna(100.0)) / 2.0)
            self._generate_threshold_variations(
                df,
                resource_scores,
                thresholds,
                direction="lte",
                title_prefix="Resource constraint",
                formula="(CPU% + RAM%) / 2 <= Threshold",
                out_dir=os.path.join(variations_root, "resource"),
                report_dir=os.path.join("reports", "variations", "resource"),
            )

        # Performance (normalized dimension, %)
        if "dim_performance" in df.columns:
            perf_scores = df["dim_performance"].fillna(0.0) * 100.0
            self._generate_threshold_variations(
                df,
                perf_scores,
                thresholds,
                direction="gte",
                title_prefix="Performance",
                formula="dim_performance >= Threshold",
                out_dir=os.path.join(variations_root, "performance"),
                report_dir=os.path.join("reports", "variations", "performance"),
            )

        # Explainability (normalized dimension, %)
        if "dim_explainability" in df.columns:
            expl_scores = df["dim_explainability"].fillna(0.0) * 100.0
            self._generate_threshold_variations(
                df,
                expl_scores,
                thresholds,
                direction="gte",
                title_prefix="Explainability",
                formula="dim_explainability >= Threshold",
                out_dir=os.path.join(variations_root, "explainability"),
                report_dir=os.path.join("reports", "variations", "explainability"),
            )

    def normalize_matrix(self, df: pd.DataFrame, criteria_info: List[Dict]) -> pd.DataFrame:
        norm_df = df.copy()
        for crit in criteria_info:
            key = crit['key']
            if key not in df.columns: continue
            values = df[key].to_numpy().astype(float)
            if crit['type'] == 'benefit':
                max_val = np.max(values)
                norm_df[key] = values / (max_val + 1e-9)
            else:
                min_val = np.min(values)
                norm_df[key] = min_val / (values + 1e-9)
        return norm_df

    def generate_final_report(self, results_df: pd.DataFrame) -> str:
        df = self.run_moo_phase(results_df)
        report = ["# DECISION SUPPORT SYSTEM REPORT (DSS)"]
        for p in ["A", "B"]:
            df = self.run_mcdm_phase(df, p)
            analysis = self.run_mcda_analysis(df, p)
            report.append(f"\n## PROFILE RESULT {p}")
            report.append(f"- **Winner**: **{analysis['winner']}**")
            report.append(f"- **Consensus**: {analysis['spearman_corr']:.4f}")
        return "\n".join(report)
