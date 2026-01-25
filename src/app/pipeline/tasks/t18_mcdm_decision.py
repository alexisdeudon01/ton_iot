import os
import pandas as pd
import json
import time
import yaml
import joblib
import numpy as np
import polars as pl
from sklearn.metrics import f1_score, recall_score, roc_auc_score
from src.mcdm.metrics_utils import compute_f_perf, compute_f_expl
from src.core.dag.task import Task
from src.core.dag.context import DAGContext
from src.core.dag.result import TaskResult
from src.app.pipeline.registry import TaskRegistry
from mcdm.decision_agent import DDoSDecisionAgent

@TaskRegistry.register("T18_MCDM_Decision")
class T18_MCDM_Decision(Task):
    """
    Final pipeline task: runs MCDM/MOO analysis to select the best algorithm.
    Updates the JSON report with decision outputs and links to charts.
    """
    def run(self, context: DAGContext) -> TaskResult:
        start_ts = time.time()
        cfg = context.config
        output_dir = "reports"
        os.makedirs(output_dir, exist_ok=True)
        
        # 1. Load evaluation results (T17)
        report_path = os.path.join("reports", "run_report.json")
        if not os.path.exists(report_path):
            return TaskResult(task_name=self.name, status="failed", duration_s=time.time()-start_ts, error="Evaluation report not found.")
            
        with open(report_path, "r") as f:
            all_metrics = json.load(f)
            
        # 2. Build decision matrix
        rows = []
        for model_name, metrics in all_metrics.items():
            if model_name == "_metadata": continue
            if not model_name.startswith("fused_"):
                continue
            if model_name == "fused_global":
                continue
            
            mcdm_inputs = metrics.get("mcdm_inputs", {})
            shap_std = mcdm_inputs.get("shap_std", 0.5)
            n_params_raw = mcdm_inputs.get("n_params") or metrics.get("complexity") or metrics.get("n_params")
            try:
                n_params = int(n_params_raw) if n_params_raw is not None else 1
            except Exception:
                n_params = 1
            n_params = max(n_params, 1)
            row = {
                "model": model_name,
                "f1": metrics.get("f1", 0),
                "recall": metrics.get("recall", 0),
                "precision": metrics.get("precision", 0),
                "accuracy": metrics.get("accuracy", 0),
                "faithfulness": metrics.get("faithfulness", mcdm_inputs.get("s_intrinsic", 0.0)),
                "stability": metrics.get("stability", max(0.0, 1.0 - shap_std)),
                "complexity": n_params,
                "latency": metrics.get("latency", 1.0),
                "cpu_percent": metrics.get("cpu_percent", 0.0),
                "ram_percent": metrics.get("ram_percent", 0.0)
            }
            rows.append(row)
            
        df_results = pd.DataFrame(rows)
        
        # 3. Initialize decision agent
        with open("config/pipeline.yaml", "r") as f:
            mcdm_config = yaml.safe_load(f)
        agent = DDoSDecisionAgent(mcdm_config, logger=context.logger)
        
        # 4. Run analysis
        ranked_df = agent.rank_models(df_results)
        
        # 5. Generate report and charts
        report = agent.generate_final_report(ranked_df)
        report_md_path = os.path.join(output_dir, "final_justification_report.md")
        with open(report_md_path, "w") as f:
            f.write(report)
            
        plots_dir = os.path.join("graph", "decision")
        os.makedirs(plots_dir, exist_ok=True)
        agent.visualize_sad(ranked_df, plots_dir)

        dtreeviz_dir = os.path.join("graph", "algorithms", "dtreeviz")
        dtreeviz_files = self._generate_dtreeviz(context, dtreeviz_dir)

        sampling_summary = self._generate_sampling_variation(context, all_metrics)
        report_links = self._write_graphs_report(sampling_summary)

        # 6. Update JSON report with MCDM outputs and chart links
        if "_metadata" in all_metrics:
            # Add decision files
            decision_files = []
            for root, _, files in os.walk(plots_dir):
                for f in files:
                    if not f.endswith(".png"):
                        continue
                    path = os.path.abspath(os.path.join(root, f))
                    decision_files.append({
                        "name": f,
                        "path": path,
                        "url": f"file://{path}",
                        "type": "decision_plot"
                    })
            algo_viz_files = [
                {
                    "name": os.path.basename(p),
                    "path": os.path.abspath(p),
                    "url": f"file://{os.path.abspath(p)}",
                    "type": "algorithm_viz"
                }
                for p in dtreeviz_files
            ]
            variation_reports = []
            variations_dir = os.path.join("reports", "variations")
            if os.path.exists(variations_dir):
                for root, _, files in os.walk(variations_dir):
                    for f in files:
                        if not (f.endswith(".docx") or f.endswith(".md")):
                            continue
                        path = os.path.abspath(os.path.join(root, f))
                        variation_reports.append({
                            "name": f,
                            "path": path,
                            "url": f"file://{path}",
                            "type": "variation_report"
                        })
            all_metrics["_metadata"]["outputs"]["generated_files"].extend(decision_files)
            all_metrics["_metadata"]["outputs"]["generated_files"].extend(algo_viz_files)
            all_metrics["_metadata"]["outputs"]["generated_files"].extend(variation_reports)
            for entry in report_links:
                all_metrics["_metadata"]["outputs"]["generated_files"].append(entry)
            all_metrics["_metadata"]["outputs"]["final_report_markdown"] = {
                "path": os.path.abspath(report_md_path),
                "url": f"file://{os.path.abspath(report_md_path)}"
            }
            
            with open(report_path, "w") as f:
                json.dump(all_metrics, f, indent=4)
        
        # Optional: generate a consolidated pipeline document with charts
        pipeline_doc = self._generate_pipeline_report(context, sampling_summary)
        if pipeline_doc and "_metadata" in all_metrics:
            all_metrics["_metadata"]["outputs"]["generated_files"].append({
                "name": os.path.basename(pipeline_doc),
                "path": os.path.abspath(pipeline_doc),
                "url": f"file://{os.path.abspath(pipeline_doc)}",
                "type": "pipeline_report",
            })
            with open(report_path, "w") as f:
                json.dump(all_metrics, f, indent=4)

        context.logger.info("mcdm", f"MCDM analysis completed. Report: {report_md_path}")
        
        return TaskResult(
            task_name=self.name, 
            status="ok", 
            duration_s=time.time() - start_ts,
            outputs=[report_md_path, plots_dir],
            meta={"winner": ranked_df.iloc[0]['model']}
        )

    def _generate_dtreeviz(self, context: DAGContext, output_dir: str) -> list:
        os.makedirs(output_dir, exist_ok=True)
        try:
            try:
                from dtreeviz import model as dtreeviz_model
            except Exception:
                from dtreeviz.trees import dtreeviz as dtreeviz_model
        except Exception as exc:
            context.logger.warning("mcdm", f"dtreeviz not available: {exc}")
            return []
        try:
            import warnings
            warnings.filterwarnings("ignore", message="findfont: Font family 'Arial' not found.")
        except Exception:
            pass

        files = []
        dataset_specs = [
            ("cic", "cic_projected", "preprocess_cic.joblib", "cic_DT.model", "DT"),
            ("ton", "ton_projected", "preprocess_ton.joblib", "ton_DT.model", "DT"),
            ("cic", "cic_projected", "preprocess_cic.joblib", "cic_RF.model", "RF"),
            ("ton", "ton_projected", "preprocess_ton.joblib", "ton_RF.model", "RF"),
        ]

        for dataset, table_id, prep_name, model_name, algo in dataset_specs:
            model_path = os.path.join("work", "models", model_name)
            prep_path = os.path.join("work", "artifacts", prep_name)

            if not os.path.exists(model_path):
                context.logger.warning("mcdm", f"Model missing for dtreeviz: {model_path}")
                continue

            try:
                table_art = context.artifact_store.load_table(table_id)
            except Exception as exc:
                context.logger.warning("mcdm", f"Table missing for dtreeviz: {table_id} ({exc})")
                continue

            if not os.path.exists(table_art.path):
                context.logger.warning("mcdm", f"Table path missing: {table_art.path}")
                continue

            df = context.table_io.read_parquet(table_art.path).collect().to_pandas()
            if df.empty:
                context.logger.warning("mcdm", f"Empty table for dtreeviz: {table_id}")
                continue

            if "y" in df.columns:
                y = df["y"].fillna(0).astype(int)
            else:
                y = pd.Series(np.zeros(len(df), dtype=int))

            if len(df) > 500:
                df = df.sample(n=500, random_state=42)
                y = y.loc[df.index]

            model_bundle = joblib.load(model_path)
            model = model_bundle.get("model")
            feature_order = model_bundle.get("feature_order", [c for c in df.columns if c != "y"])

            X_raw = df[feature_order]
            feature_names = feature_order
            if os.path.exists(prep_path):
                try:
                    preprocessor = joblib.load(prep_path)
                    X = preprocessor.transform(X_raw)
                    try:
                        feature_names = list(preprocessor.get_feature_names_out(feature_order))
                    except Exception:
                        feature_names = feature_order
                except Exception as exc:
                    context.logger.warning("mcdm", f"Preprocess load failed for {prep_path}: {exc}")
                    X = X_raw.values
            else:
                X = X_raw.values

            tree_model = model
            if algo == "RF":
                try:
                    tree_model = model.estimators_[0]
                except Exception:
                    context.logger.warning("mcdm", f"No estimator available in RF for {model_path}")
                    continue

            class_names = [str(c) for c in sorted(y.unique())]
            try:
                viz = dtreeviz_model(
                    tree_model,
                    X,
                    y.values,
                    target_name="label",
                    feature_names=feature_names,
                    class_names=class_names
                )
                out_name = f"{dataset.lower()}_{algo.lower()}_dtreeviz.svg"
                out_path = os.path.join(output_dir, out_name)
                render = viz.view()
                if hasattr(render, "save"):
                    render.save(out_path)
                elif hasattr(render, "save_svg"):
                    svg = render.save_svg()
                    with open(out_path, "w") as f:
                        f.write(svg)
                else:
                    with open(out_path, "w") as f:
                        f.write(render.svg())
                files.append(out_path)
            except Exception as exc:
                context.logger.warning("mcdm", f"dtreeviz failed for {model_path}: {exc}")
                continue

        return files

    def _generate_sampling_variation(self, context: DAGContext, all_metrics: dict) -> dict:
        output_root = os.path.join("graph", "decision", "variations", "sampling")
        perf_dir = os.path.join(output_root, "performance")
        expl_dir = os.path.join(output_root, "explainability")
        os.makedirs(perf_dir, exist_ok=True)
        os.makedirs(expl_dir, exist_ok=True)

        pred_cic_path = os.path.join("work", "data", "predictions_cic.parquet")
        pred_ton_path = os.path.join("work", "data", "predictions_ton.parquet")
        if not (os.path.exists(pred_cic_path) and os.path.exists(pred_ton_path)):
            context.logger.warning("mcdm", "Sampling variation skipped: prediction files missing")
            return {"files": [], "perf_derivative_peaks": {}, "expl_derivative_peaks": {}}

        df_cic = pl.read_parquet(pred_cic_path).filter(pl.col("split") == "test")
        df_ton = pl.read_parquet(pred_ton_path).filter(pl.col("split") == "test")
        if df_cic.is_empty() or df_ton.is_empty():
            context.logger.warning("mcdm", "Sampling variation skipped: empty test split")
            return {"files": [], "perf_derivative_peaks": {}, "expl_derivative_peaks": {}}

        models = sorted(set(df_cic["model"].unique().to_list()) & set(df_ton["model"].unique().to_list()))
        if not models:
            context.logger.warning("mcdm", "Sampling variation skipped: no common models")
            return {"files": [], "perf_derivative_peaks": {}, "expl_derivative_peaks": {}}

        fractions = [i / 10 for i in range(1, 11)]
        context.logger.info(
            "mcdm",
            "Sampling variation: evaluating fractions",
            fractions=[int(f * 100) for f in fractions],
            models=models,
        )
        results_perf = {m: [] for m in models}
        results_expl = {m: [] for m in models}
        stds = {m: [] for m in models}
        logged_base = set()

        fused_metrics = {k.replace("fused_", ""): v for k, v in all_metrics.items() if k.startswith("fused_") and k != "fused_global"}

        for frac in fractions:
            for model in models:
                df_cic_m = df_cic.filter(pl.col("model") == model)
                df_ton_m = df_ton.filter(pl.col("model") == model)
                if df_cic_m.is_empty() or df_ton_m.is_empty():
                    results_perf[model].append(0.0)
                    results_expl[model].append(0.0)
                    stds[model].append(0.0)
                    continue

                if model not in logged_base:
                    context.logger.info(
                        "mcdm",
                        "Sampling base counts",
                        model=model,
                        cic_count=len(df_cic_m),
                        ton_count=len(df_ton_m),
                    )
                    logged_base.add(model)

                n_cic = max(1, int(len(df_cic_m) * frac))
                n_ton = max(1, int(len(df_ton_m) * frac))
                context.logger.info(
                    "mcdm",
                    "Sampling fraction",
                    model=model,
                    fraction=frac,
                    n_cic=n_cic,
                    n_ton=n_ton,
                )
                sample_cic = df_cic_m.sample(n=n_cic, seed=int(42 + frac * 100)).to_pandas()
                sample_ton = df_ton_m.sample(n=n_ton, seed=int(84 + frac * 100)).to_pandas()

                def _safe_auc(y_true, y_score):
                    try:
                        if len(set(y_true)) < 2:
                            return 0.5
                        return float(roc_auc_score(y_true, y_score))
                    except Exception:
                        return 0.5

                def _compute_basic(sample):
                    y_true = sample["y_true"].to_numpy()
                    proba = sample["proba"].to_numpy()
                    y_pred = (proba >= 0.5).astype(int)
                    f1 = float(f1_score(y_true, y_pred, zero_division=0))
                    recall = float(recall_score(y_true, y_pred, zero_division=0))
                    auc = _safe_auc(y_true, proba)
                    return f1, recall, auc, y_true, proba

                f1_cic, recall_cic, auc_cic, y_cic, p_cic = _compute_basic(sample_cic)
                f1_ton, recall_ton, auc_ton, y_ton, p_ton = _compute_basic(sample_ton)

                y_combined = np.concatenate([y_cic, y_ton])
                p_combined = np.concatenate([p_cic, p_ton])
                y_pred_combined = (p_combined >= 0.5).astype(int)
                f1_combined = float(f1_score(y_combined, y_pred_combined, zero_division=0))
                recall_combined = float(recall_score(y_combined, y_pred_combined, zero_division=0))
                auc_combined = _safe_auc(y_combined, p_combined)
                gap = abs(f1_cic - f1_ton)
                f_perf = float(compute_f_perf(f1_combined, recall_combined, auc_combined, gap))
                results_perf[model].append(f_perf)
                stds[model].append(float(np.std(p_combined)))

        for model in models:
            max_std = max(stds[model]) if stds[model] else 0.0
            if max_std <= 0:
                max_std = 1.0
            metrics = fused_metrics.get(model, {})
            mcdm_inputs = metrics.get("mcdm_inputs", {})
            s_intrinsic = float(metrics.get("faithfulness", mcdm_inputs.get("s_intrinsic", 0.0)))
            shap_available = bool(mcdm_inputs.get("shap_available", False))
            n_params_raw = mcdm_inputs.get("n_params") or metrics.get("complexity") or metrics.get("n_params")
            try:
                n_params = int(n_params_raw) if n_params_raw is not None else 1
            except Exception:
                n_params = 1
            n_params = max(n_params, 1)
            for std_val in stds[model]:
                stability = max(0.0, 1.0 - (std_val / max_std))
                shap_std = 1.0 - stability
                f_expl = float(compute_f_expl(s_intrinsic, shap_available, n_params, shap_std))
                results_expl[model].append(f_expl)

        x_vals = [int(frac * 100) for frac in fractions]
        files = []

        def _plot_lines(data_map, title, y_label, out_path):
            import matplotlib.pyplot as plt
            plt.figure(figsize=(10, 6))
            for model, values in data_map.items():
                plt.plot(x_vals, values, marker='o', label=model)
            plt.title(title)
            plt.xlabel("Sampling (%)")
            plt.ylabel(y_label)
            plt.grid(alpha=0.3)
            plt.legend()
            plt.tight_layout()
            plt.savefig(out_path, dpi=150, bbox_inches='tight')
            plt.close()

        perf_plot = os.path.join(perf_dir, "performance_sampling_curve.png")
        _plot_lines(results_perf, "Performance vs Sampling", "f_perf", perf_plot)
        files.append(perf_plot)

        expl_plot = os.path.join(expl_dir, "explainability_sampling_curve.png")
        _plot_lines(results_expl, "Explainability vs Sampling", "f_expl", expl_plot)
        files.append(expl_plot)

        perf_derivative_peaks = {}
        expl_derivative_peaks = {}

        def _plot_derivative(data_map, title, y_label, out_path, peaks_out):
            import matplotlib.pyplot as plt
            plt.figure(figsize=(10, 6))
            for model, values in data_map.items():
                deriv = np.gradient(values, x_vals)
                plt.plot(x_vals, deriv, marker='o', label=model)
                idx = int(np.argmax(np.abs(deriv)))
                peaks_out[model] = {"peak": float(deriv[idx]), "at": x_vals[idx]}
            plt.axhline(0, color='black', linewidth=1)
            plt.title(title)
            plt.xlabel("Sampling (%)")
            plt.ylabel(y_label)
            plt.grid(alpha=0.3)
            plt.legend()
            plt.tight_layout()
            plt.savefig(out_path, dpi=150, bbox_inches='tight')
            plt.close()

        perf_deriv_plot = os.path.join(perf_dir, "performance_sampling_derivative.png")
        _plot_derivative(results_perf, "Performance derivative vs Sampling", "d(f_perf)/d(sampling)", perf_deriv_plot, perf_derivative_peaks)
        files.append(perf_deriv_plot)

        expl_deriv_plot = os.path.join(expl_dir, "explainability_sampling_derivative.png")
        _plot_derivative(results_expl, "Explainability derivative vs Sampling", "d(f_expl)/d(sampling)", expl_deriv_plot, expl_derivative_peaks)
        files.append(expl_deriv_plot)

        return {
            "files": files,
            "perf_derivative_peaks": perf_derivative_peaks,
            "expl_derivative_peaks": expl_derivative_peaks,
        }

    def _write_graphs_report(self, sampling_summary: dict) -> list:
        report_path = os.path.join("reports", "analysis_graphs_report.md")
        os.makedirs(os.path.dirname(report_path), exist_ok=True)

        def _collect_files(root, exts):
            collected = []
            if not os.path.exists(root):
                return collected
            for dirpath, _, files in os.walk(root):
                for f in files:
                    if any(f.endswith(ext) for ext in exts):
                        rel = os.path.join(dirpath, f)
                        collected.append(rel)
            return sorted(collected)

        sections = [
            ("Decision charts", "graph/decision", [".png"]),
            ("Threshold variations", "graph/decision/variations", [".png"]),
            ("Feature distributions", "graph/feature_distributions", [".png", ".md"]),
            ("Dtreeviz", "graph/algorithms/dtreeviz", [".svg", ".png"]),
        ]

        lines = ["# Charts report", ""]
        for title, root, exts in sections:
            files = _collect_files(root, exts)
            lines.append(f"## {title}")
            if not files:
                lines.append("- No files.")
                lines.append("")
                continue
            for f in files:
                lines.append(f"- {f}")
            lines.append("")

        lines.append("## Derivative summary (sampling)")
        if sampling_summary.get("perf_derivative_peaks"):
            lines.append("### Performance")
            for model, info in sampling_summary["perf_derivative_peaks"].items():
                lines.append(f"- {model}: derivative peak {info['peak']:.4f} at {info['at']}%")
        else:
            lines.append("- No performance derivative data.")
        if sampling_summary.get("expl_derivative_peaks"):
            lines.append("### Explainability")
            for model, info in sampling_summary["expl_derivative_peaks"].items():
                lines.append(f"- {model}: derivative peak {info['peak']:.4f} at {info['at']}%")
        else:
            lines.append("- No explainability derivative data.")
        lines.append("")

        with open(report_path, "w") as f:
            f.write("\n".join(lines))

        return [{
            "name": os.path.basename(report_path),
            "path": os.path.abspath(report_path),
            "url": f"file://{os.path.abspath(report_path)}",
            "type": "graphs_report"
        }]

    def _generate_pipeline_report(self, context: DAGContext, sampling_summary: dict) -> str:
        """Create a consolidated DOCX/MD report with key pipeline graphs and explanations."""
        output_path = os.path.join("work", "pipeline_visual_report.docx")
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        def _collect_pngs(root, recursive=True):
            if not os.path.exists(root):
                return []
            files = []
            if recursive:
                for dirpath, _, filenames in os.walk(root):
                    for name in sorted(filenames):
                        if name.endswith(".png"):
                            files.append(os.path.join(dirpath, name))
            else:
                for name in sorted(os.listdir(root)):
                    if name.endswith(".png"):
                        files.append(os.path.join(root, name))
            return files

        sections = [
            ("Pipeline overview", [os.path.join("topsis_tppis", "09_pipeline_overview.png")]),
            ("Stratification checks", _collect_pngs(os.path.join("graph", "stratification"))),
            ("Feature distributions", _collect_pngs(os.path.join("graph", "feature_distributions"))),
            ("Decision charts (MCDM)", _collect_pngs(os.path.join("graph", "decision"), recursive=False)),
            ("Threshold variations", _collect_pngs(os.path.join("graph", "decision", "variations"))),
            ("TOPSIS visual report", _collect_pngs(os.path.join("topsis_tppis"))),
        ]
        section_notes = {
            "Pipeline overview": "High-level flow of the MCDA pipeline and where each visualization fits.",
            "Stratification checks": "Compares label proportions before and after sampling to verify stratification.",
            "Feature distributions": "Shows post-transformation feature distributions used by the models.",
            "Decision charts (MCDM)": "Summarizes decision matrix, TOPSIS steps, and ranking views.",
            "Threshold variations": "Shows admissible vs rejected solutions across threshold levels.",
            "TOPSIS visual report": "Includes KS validation, feature importance, decision matrices, and final ranking.",
        }

        def _describe_file(name: str) -> str:
            key = name.lower()
            if "stratification" in key:
                return "Bar chart comparing full vs sampled label proportions."
            if "feature_importance" in key:
                return "Feature ranking by relevance to the target label."
            if "ks_validation" in key:
                return "KS statistics per feature against the theoretical threshold."
            if "raw_matrix" in key:
                return "Raw decision matrix heatmap (before normalization)."
            if "normalized_matrix" in key:
                return "Normalized decision matrix heatmap."
            if "ahp" in key or "weights" in key:
                return "AHP weights used for TOPSIS."
            if "topsis" in key or "closeness" in key:
                return "TOPSIS closeness coefficients for each alternative."
            if "radar" in key:
                return "Radar comparison of top alternatives."
            if "pareto" in key:
                return "Pareto front visualization of non-dominated vs dominated solutions."
            if "sampling" in key and "derivative" in key:
                return "Derivative plot showing sensitivity of the metric to sampling rate."
            if "sampling" in key:
                return "Metric evolution across sampling ratios."
            if "pipeline_overview" in key:
                return "Pipeline flow diagram for MCDA steps."
            return "Chart generated by the pipeline."

        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()
            doc.add_heading("Pipeline Visual Report", level=0)
            doc.add_paragraph(
                "This document consolidates the pipeline outputs with key graphs and short interpretations."
            )

            doc.add_heading("Sampling variation summary", level=1)
            if sampling_summary.get("perf_derivative_peaks"):
                doc.add_paragraph("Performance derivative peaks:")
                for model, info in sampling_summary["perf_derivative_peaks"].items():
                    doc.add_paragraph(
                        f"{model}: peak {info.get('peak', 0):.4f} at {info.get('at', 0)}%",
                        style="List Bullet",
                    )
            if sampling_summary.get("expl_derivative_peaks"):
                doc.add_paragraph("Explainability derivative peaks:")
                for model, info in sampling_summary["expl_derivative_peaks"].items():
                    doc.add_paragraph(
                        f"{model}: peak {info.get('peak', 0):.4f} at {info.get('at', 0)}%",
                        style="List Bullet",
                    )
            doc.add_paragraph("Threshold variation range: 100 to 30 (step 10) for resources, performance, and explainability.")

            for title, files in sections:
                if not files:
                    continue
                doc.add_heading(title, level=1)
                doc.add_paragraph(section_notes.get(title, "The following charts document this pipeline stage."))
                for path in files:
                    if not os.path.exists(path):
                        continue
                    doc.add_paragraph(os.path.basename(path))
                    doc.add_paragraph(_describe_file(os.path.basename(path)))
                    doc.add_picture(path, width=Inches(6.5))

            doc.save(output_path)
            context.logger.info("mcdm", f"Pipeline visual report saved: {output_path}")
            return output_path
        except Exception as exc:
            md_path = output_path.replace(".docx", ".md")
            lines = ["# Pipeline Visual Report", ""]
            lines.append("This document consolidates the pipeline outputs with key graphs.")
            lines.append("")
            lines.append("## Sampling variation summary")
            if sampling_summary.get("perf_derivative_peaks"):
                lines.append("### Performance derivative peaks")
                for model, info in sampling_summary["perf_derivative_peaks"].items():
                    lines.append(f"- {model}: peak {info.get('peak', 0):.4f} at {info.get('at', 0)}%")
            if sampling_summary.get("expl_derivative_peaks"):
                lines.append("### Explainability derivative peaks")
                for model, info in sampling_summary["expl_derivative_peaks"].items():
                    lines.append(f"- {model}: peak {info.get('peak', 0):.4f} at {info.get('at', 0)}%")
            lines.append("- Threshold variation range: 100 to 30 (step 10) for resources, performance, and explainability.")
            for title, files in sections:
                if not files:
                    continue
                lines.append("")
                lines.append(f"## {title}")
                note = section_notes.get(title)
                if note:
                    lines.append(note)
                for path in files:
                    if os.path.exists(path):
                        lines.append(f"- {path}  ({_describe_file(os.path.basename(path))})")
            with open(md_path, "w") as f:
                f.write("\n".join(lines))
            context.logger.warning("mcdm", f"Docx unavailable; wrote markdown report: {md_path}. Error: {exc}")
            return md_path
